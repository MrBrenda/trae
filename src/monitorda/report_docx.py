"""DOCX 报告：复现 0423 报告核心表格结构。

策略：手工构造章节 + 表格。后续可改为加载 templates/template.docx 套样式。
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt

from .io_paths import paths


def render_docx(
    run_dir: Path,
    *,
    run_date: date | None = None,
    window_start: str | None = None,
    window_end: str | None = None,
) -> Path:
    p = paths()
    run_dir.mkdir(parents=True, exist_ok=True)

    events_df = _safe_read(p.parquet("events"))
    diag_df = _safe_read(p.parquet("node_diagnostics"))
    rdii_df = _safe_read(p.parquet("rdii_by_event_node"))

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "PingFang SC"
    style.font.size = Pt(10.5)

    title = doc.add_heading(f"排水管网监测数据诊断报告 — {(run_date or date.today()).isoformat()}", level=0)
    title.alignment = 1
    doc.add_paragraph(
        f"方法论锚定《红旗东路积涝整治及管网改造专题研究阶段性成果报告 0423》。"
        f"时间窗：{window_start or '(全量)'} → {window_end or '(全量)'}；"
        f"节点 {len(diag_df)} 个，事件 {len(events_df)} 场。"
    )

    # 2. 降雨事件
    doc.add_heading("一、降雨事件汇总", level=1)
    if events_df.empty:
        doc.add_paragraph("（未识别到事件）")
    else:
        _add_table(doc, events_df[[
            "event_id", "station_id", "t_start", "t_end",
            "duration_h", "total_mm", "max_intensity_mmh", "antecedent_dry_d",
        ]], header_zh=["事件 ID", "站点", "起", "止", "时长 h", "累计 mm", "峰值 mm/h", "前置干天 d"])

    # 4.4 RDII 分级
    doc.add_heading("二、节点 RDII 分级（表 4.4）", level=1)
    if rdii_df.empty:
        doc.add_paragraph("（无 RDII 计算结果）")
    else:
        # 取每节点 median
        grades = (rdii_df.groupby("node_id")
                  .agg(mean_rise=("rise_amp_m", "median"),
                       mean_lag=("lag_start_h", "median"),
                       grade=("grade", lambda s: _mode(s)))
                  .reset_index())
        _add_table(doc, grades, header_zh=["节点", "中位升幅 m", "中位时滞 h", "分级"])

    # 5.4 综合诊断
    doc.add_heading("三、节点综合诊断（表 5.4）", level=1)
    if diag_df.empty:
        doc.add_paragraph("（无诊断结果）")
    else:
        cols = ["node_id", "name_zh", "n_events", "mean_illicit_area_km2",
                "mean_rise_amp_m", "median_lag_start_h", "median_halflife_h",
                "category", "evidence_score", "notes"]
        _add_table(doc, diag_df[cols],
                   header_zh=["节点", "名称", "事件数", "平均混接面积 km²",
                              "中位升幅 m", "中位时滞 h", "半衰期 h",
                              "分类", "置信", "备注"])

    # 图
    doc.add_heading("四、图表", level=1)
    figs = [
        ("监测节点空间分布与诊断类别", run_dir / "figures" / "site_map.png"),
        ("各节点雨天响应升幅", run_dir / "figures" / "node_rise_amp.png"),
        ("第五污水厂入流时间序列", run_dir / "figures" / "plant_inlet.png"),
    ]
    for caption, fp in figs:
        if fp.exists():
            doc.add_heading(caption, level=2)
            doc.add_picture(str(fp), width=Inches(6.0))

    out = run_dir / "report.docx"
    doc.save(str(out))
    return out


def _add_table(doc: Document, df: pd.DataFrame, *, header_zh: list[str]) -> None:
    t = doc.add_table(rows=1, cols=len(header_zh))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(header_zh):
        hdr[i].text = h
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for i, col in enumerate(df.columns):
            v = row[col]
            if isinstance(v, float):
                cells[i].text = f"{v:.3f}" if abs(v) < 100 else f"{v:.1f}"
            elif v is None or (isinstance(v, float) and pd.isna(v)):
                cells[i].text = "—"
            else:
                cells[i].text = str(v)


def _mode(s: pd.Series) -> str:
    vc = s.dropna().value_counts()
    return str(vc.index[0]) if not vc.empty else "NA"


def _safe_read(p: Path) -> pd.DataFrame:
    return pd.read_parquet(p) if p.exists() else pd.DataFrame()
